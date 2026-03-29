from fastapi import FastAPI, HTTPException, UploadFile, File, Query, Depends
from pydantic import BaseModel
from typing import Optional, Dict, Any, List
import os
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, PlainTextResponse

try:
    import anthropic
except Exception:
    anthropic = None

from .parser import parse_resume
from .rag_engine import search_programs
from .cv_generator import generate_cv
from .auth import router as auth_router, users_router
from .programs import router as programs_router
from .parser import extract_profile
from .rag_engine import DATA_PATH as PROGRAMS_PATH
import json
import httpx
from .db import get_db_docs, load_db, select_relevant_chunks, DB_EXCLUDE_NAMES, KNOWLEDGE_DIR, list_db_docs
from .recommend_prompt import build_recommendation_prompt
from .schemas import validate_llm_json
from .models import seed_universities_from_json, University
from .auth import get_current_user
from .database import SessionLocal
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
def _startup_seed():
    try:
        seed_universities_from_json(PROGRAMS_PATH)
    except Exception:
        pass
class ChatInput(BaseModel):
    text: str
    system: Optional[str] = None
    model: Optional[str] = None
    max_tokens: Optional[int] = 512


class ParseInput(BaseModel):
    text: str


class RAGQuery(BaseModel):
    query: str


class CVRequest(BaseModel):
    profile: Dict[str, Any]
    target_program: Optional[str] = None


class RecommendRequest(BaseModel):
    profile: Dict[str, Any]
    top_k: Optional[int] = 5
    locale: Optional[str] = "zh"


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/api/parse")
def api_parse(payload: ParseInput):
    return parse_resume(payload.text)


@app.post("/api/rag")
def api_rag(payload: RAGQuery):
    return {"matches": search_programs(payload.query)}


@app.post("/api/cv")
def api_cv(payload: CVRequest):
    return {"cv": generate_cv(payload.profile, payload.target_program)}


@app.post("/api/llm/chat")
def api_chat(payload: ChatInput):
    api_key = os.getenv("ANTHROPIC_API_KEY")
    base_url = os.getenv("ANTHROPIC_BASE_URL")
    if anthropic is None:
        raise HTTPException(status_code=500, detail="anthropic sdk not installed")
    if not api_key:
        raise HTTPException(status_code=400, detail="missing ANTHROPIC_API_KEY")
    client = anthropic.Anthropic(api_key=api_key, base_url=base_url) if hasattr(anthropic, "Anthropic") else None
    if client is None:
        raise HTTPException(status_code=500, detail="anthropic client unavailable")
    model = payload.model or "MiniMax-M2.7"
    max_tokens = payload.max_tokens or 512
    system = payload.system or "You are a helpful assistant."
    try:
        resp = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": [{"type": "text", "text": payload.text}]}],
        )
        texts: List[str] = []
        for block in resp.content:
            if getattr(block, "type", None) == "text":
                texts.append(block.text)
        return {"text": "\n".join(texts)}
    except Exception as e:
        raise HTTPException(status_code=502, detail=str(e))


@app.post("/api/cv/extract")
async def api_cv_extract(file: UploadFile = File(...)):
    content = await file.read()
    text = ""
    try:
        fname = (file.filename or "").lower()
        if fname.endswith(".pdf"):
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                t = page.extract_text() or ""
                if t:
                    text += t
            if not text.strip():
                try:
                    from pdfminer.high_level import extract_text as pdfminer_extract_text
                    import io
                    text = pdfminer_extract_text(io.BytesIO(content)) or ""
                except Exception:
                    pass
        elif fname.endswith(".docx"):
            try:
                import io
                from docx import Document
                doc = Document(io.BytesIO(content))
                text = "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                text = content.decode("utf-8", errors="ignore")
        else:
            text = content.decode("utf-8", errors="ignore")
    except Exception:
        text = content.decode("utf-8", errors="ignore")
    if (file.filename or "").lower().endswith(".pdf") and not text.strip() and os.getenv("ENABLE_OCR", "").lower() in ("1", "true", "yes"):
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            images = convert_from_bytes(content, dpi=200)
            ocr_text = ""
            for img in images:
                ocr_text += pytesseract.image_to_string(img)
            if ocr_text.strip():
                text = ocr_text
        except Exception:
            pass
    profile = extract_profile(text)
    return {"profile": profile, "meta": {"text_chars": len(text), "ocr_used": bool(text and os.getenv("ENABLE_OCR", "").lower() in ("1", "true", "yes"))}}

@app.post("/api/cv/extract_batch")
async def api_cv_extract_batch(files: List[UploadFile] = File(...)):
    results: List[Dict[str, Any]] = []
    for file in files:
        try:
            content = await file.read()
            text = ""
            try:
                fname = (file.filename or "").lower()
                if fname.endswith(".pdf"):
                    from pypdf import PdfReader
                    import io
                    reader = PdfReader(io.BytesIO(content))
                    for page in reader.pages:
                        t = page.extract_text() or ""
                        if t:
                            text += t
                    if not text.strip():
                        try:
                            from pdfminer.high_level import extract_text as pdfminer_extract_text
                            import io
                            text = pdfminer_extract_text(io.BytesIO(content)) or ""
                        except Exception:
                            pass
                elif fname.endswith(".docx"):
                    try:
                        import io
                        from docx import Document
                        doc = Document(io.BytesIO(content))
                        text = "\n".join(p.text for p in doc.paragraphs)
                    except Exception:
                        text = content.decode("utf-8", errors="ignore")
                else:
                    text = content.decode("utf-8", errors="ignore")
            except Exception:
                text = content.decode("utf-8", errors="ignore")
            if (file.filename or "").lower().endswith(".pdf") and not text.strip() and os.getenv("ENABLE_OCR", "").lower() in ("1", "true", "yes"):
                try:
                    from pdf2image import convert_from_bytes
                    import pytesseract
                    images = convert_from_bytes(content, dpi=200)
                    ocr_text = ""
                    for img in images:
                        ocr_text += pytesseract.image_to_string(img)
                    if ocr_text.strip():
                        text = ocr_text
                except Exception:
                    pass
            profile = extract_profile(text)
            results.append({
                "file_name": file.filename,
                "profile": profile,
                "meta": {"text_chars": len(text), "ocr_used": bool(text and os.getenv("ENABLE_OCR", "").lower() in ("1", "true", "yes"))}
            })
        except Exception as e:
            results.append({
                "file_name": file.filename if file else None,
                "error": str(e)
            })
    return {"results": results}

class TextPayload(BaseModel):
    text: str

@app.post("/api/cv/extract_text")
def api_cv_extract_text(payload: TextPayload):
    profile = extract_profile(payload.text)
    return {"profile": profile}


@app.post("/api/recommend")
def api_recommend(payload: RecommendRequest):
    try:
        data = json.loads(PROGRAMS_PATH.read_text(encoding="utf-8"))
    except Exception:
        data = []
    load_db()
    prof = payload.profile or {}
    def to_float(v):
        try:
            if v in (None, "", " "):
                return None
            return float(v)
        except Exception:
            return None
    def to_int(v):
        try:
            if v in (None, "", " "):
                return None
            return int(float(v))
        except Exception:
            return None
    gpa = to_float(prof.get("gpa"))
    toefl = to_int(prof.get("toefl"))
    ielts = to_float(prof.get("ielts"))
    hsk = to_int(prof.get("hsk"))
    results: List[Dict[str, Any]] = []
    for p in data:
        req = p.get("requirements", {})
        matched = []
        missing = []
        # GPA
        gpa_min = req.get("gpa_min")
        if gpa_min is not None:
            if gpa is None:
                missing.append("GPA info missing")
            elif gpa < gpa_min:
                missing.append(f"GPA below requirement: {gpa} < {gpa_min}")
            else:
                matched.append(f"GPA meets requirement: {gpa} ≥ {gpa_min}")
        # TOEFL
        toefl_min = req.get("toefl_min")
        if toefl_min is not None:
            if toefl is None:
                missing.append("TOEFL score missing")
            elif toefl < toefl_min:
                missing.append(f"TOEFL below requirement: {toefl} < {toefl_min}")
            else:
                matched.append(f"TOEFL meets requirement: {toefl} ≥ {toefl_min}")
        # IELTS
        ielts_min = req.get("ielts_min")
        if ielts_min is not None:
            if ielts is None:
                missing.append("IELTS score missing")
            elif ielts < ielts_min:
                missing.append(f"IELTS below requirement: {ielts} < {ielts_min}")
            else:
                matched.append(f"IELTS meets requirement: {ielts} ≥ {ielts_min}")
        # HSK
        hsk_required = req.get("hsk_required")
        if hsk_required:
            if hsk is None:
                missing.append("HSK certificate missing")
            elif hsk <= 0:
                missing.append(f"Invalid HSK value: {hsk}")
            else:
                matched.append(f"HSK provided: {hsk}")
        # Special requirements
        special_req = req.get("special_requirements")
        if special_req:
            missing.append("Special requirements may apply")
        # Fees and deadline soft info
        fees = p.get("fees")
        deadline = p.get("deadline")
        if fees is not None:
            matched.append(f"Tuition per year: {fees}")
        if deadline:
            matched.append(f"Application deadline: {deadline}")
        satisfied = len(matched)
        unmet = len(missing)
        qs = p.get("qs_rank") or 1e9
        score = satisfied * 8 - unmet * 7 - (qs / 120.0)
        if isinstance(fees, (int, float)):
            score += max(0, 50000 - fees) / 5000.0
        from datetime import datetime
        try:
            if deadline:
                dt = datetime.fromisoformat(str(deadline))
                days_left = (dt - datetime.now()).days
                if days_left < 0:
                    score -= 20
                    missing.append("Deadline has passed")
                elif days_left <= 30:
                    score -= 5
        except Exception:
            pass
        if special_req:
            score -= 3
        results.append({
            "id": p.get("id"),
            "name": p.get("name"),
            "university": p.get("university"),
            "city": p.get("city"),
            "official_link": p.get("official_link"),
            "fees": fees,
            "deadline": deadline,
            "requirements": req,
            "matched_elements": matched,
            "missing_elements": missing,
            "is_eligible": unmet == 0,
            "match_score": score,
            "link": f"/api/programs/{p.get('id')}",
            "client_route": f"#detail?program={p.get('id')}",
            "qs_rank": p.get("qs_rank"),
            "cn_rank": p.get("cn_rank"),
            "fields": p.get("fields", []),
        })
    results.sort(key=lambda x: (not x["is_eligible"], -x["match_score"]))
    top = results[: (payload.top_k or 5)]
    llm_analysis = None
    llm_json = None
    minimax_key = os.getenv("MINIMAX_API_KEY")
    minimax_base = os.getenv("MINIMAX_BASE_URL", "https://api.minimax.chat/v1")
    minimax_model = os.getenv("MINIMAX_MODEL", "abab6.5-chat")
    if minimax_key:
        try:
            chunks = select_relevant_chunks(prof, limit_chars=6000)
            full_profile = {
                "gpa": gpa,
                "toefl": toefl,
                "ielts": ielts,
                "hsk": hsk,
                "hsk_level": prof.get("hsk_level"),
                "hsk_score": prof.get("hsk_score"),
                "majors": prof.get("majors", []),
                "cities": prof.get("cities", []),
                "preferred_fields": prof.get("preferred_fields", []),
                "preferred_cities": prof.get("preferred_cities", []),
                "math_physics_avg": prof.get("math_physics_avg"),
            }
            prompt = build_recommendation_prompt(full_profile, top, chunks, locale=payload.locale or "en")
            payload_llm = {
                "model": minimax_model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 800
            }
            url = f"{minimax_base.rstrip('/')}/text/chatcompletions"
            headers = {"Authorization": f"Bearer {minimax_key}", "Content-Type": "application/json"}
            resp = httpx.post(url, json=payload_llm, headers=headers, timeout=45)
            data_llm = resp.json()
            content = None
            if isinstance(data_llm, dict) and data_llm.get("choices"):
                ch0 = data_llm["choices"][0]
                msg = ch0.get("message") if isinstance(ch0, dict) else None
                if isinstance(msg, dict):
                    content = msg.get("content")
                if content is None and isinstance(ch0, dict):
                    content = ch0.get("text")
            if isinstance(content, str):
                llm_analysis = content
                try:
                    import re, json as _json
                    m = re.search(r"\{[\s\S]*\}", content)
                    if m:
                        parsed = _json.loads(m.group(0))
                        llm_json = validate_llm_json(parsed)
                        if llm_json is None:
                            from .schemas import coerce_llm_json
                            llm_json = coerce_llm_json(parsed)
                except Exception:
                    llm_json = None
        except Exception:
            llm_analysis = None
    return {"recommendations": top, "llm_analysis": llm_analysis, "llm_json": llm_json}


app.include_router(auth_router)
app.include_router(users_router)
app.include_router(programs_router)


@app.get("/api/cv/sample", response_class=PlainTextResponse)
def api_cv_sample(name: Optional[str] = Query(default=None)):
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    samples_dir = os.path.join(base_dir, "samples")
    exts = {".txt", ".pdf", ".docx"}
    candidates: List[str] = []
    if os.path.isdir(samples_dir):
        for fn in os.listdir(samples_dir):
            lower = fn.lower()
            if any(lower.endswith(e) for e in exts):
                candidates.append(os.path.join(samples_dir, fn))
    for fn in os.listdir(base_dir):
        lower = fn.lower()
        if any(lower.endswith(e) for e in exts):
            if fn in DB_EXCLUDE_NAMES:
                continue
            candidates.append(os.path.join(base_dir, fn))
    target_path = None
    if name:
        for p in candidates:
            if os.path.basename(p) == name:
                target_path = p
                break
    if target_path is None:
        default_path = os.path.join(base_dir, "CV – ARIE PRATAMA.txt")
        target_path = default_path if os.path.isfile(default_path) else (candidates[0] if candidates else None)
    if not target_path or not os.path.isfile(target_path):
        raise HTTPException(status_code=404, detail="sample CV not found")
    if target_path.lower().endswith(".pdf"):
        try:
            from pypdf import PdfReader
            import io
            with open(target_path, "rb") as rf:
                reader = PdfReader(io.BytesIO(rf.read()))
            out = ""
            for page in reader.pages:
                out += page.extract_text() or ""
            if not out.strip():
                from pdfminer.high_level import extract_text as pdfminer_extract_text
                out = pdfminer_extract_text(target_path) or ""
            return out
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    if target_path.lower().endswith(".docx"):
        try:
            from docx import Document
            doc = Document(target_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
    try:
        with open(target_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/cv/samples")
def api_cv_samples():
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    samples_dir = os.path.join(base_dir, "samples")
    exts = {".txt", ".pdf", ".docx"}
    files: List[Dict[str, Any]] = []
    def add_dir(d):
        if os.path.isdir(d):
            for fn in os.listdir(d):
                lower = fn.lower()
                if any(lower.endswith(e) for e in exts):
                    if fn in DB_EXCLUDE_NAMES:
                        continue
                    p = os.path.join(d, fn)
                    try:
                        st = os.stat(p)
                        files.append({"name": fn, "size": st.st_size})
                    except Exception:
                        files.append({"name": fn, "size": None})
    add_dir(samples_dir)
    add_dir(base_dir)
    seen = set()
    deduped = []
    for item in files:
        if item["name"] not in seen:
            seen.add(item["name"])
            deduped.append(item)
    return {"files": deduped}


@app.get("/api/db/list")
def api_db_list(current_user=Depends(get_current_user)):
    if not getattr(current_user, "is_admin", False):
        raise HTTPException(status_code=403, detail="forbidden")
    load_db()
    return {"docs": list_db_docs()}


@app.post("/api/db/upload")
async def api_db_upload(files: List[UploadFile] = File(...), current_user=Depends(get_current_user)):
    if not getattr(current_user, "is_admin", False):
        raise HTTPException(status_code=403, detail="forbidden")
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    knowledge_dir = os.path.join(base_dir, "knowledge")
    os.makedirs(knowledge_dir, exist_ok=True)
    saved = []
    for f in files:
        try:
            name = f.filename
            if not name:
                continue
            data = await f.read()
            target = os.path.join(knowledge_dir, os.path.basename(name))
            with open(target, "wb") as out:
                out.write(data)
            saved.append({"name": os.path.basename(name), "path": target})
        except Exception as e:
            saved.append({"name": f.filename, "error": str(e)})
    return {"saved": saved}


@app.post("/api/db/reload")
def api_db_reload(current_user=Depends(get_current_user)):
    if not getattr(current_user, "is_admin", False):
        raise HTTPException(status_code=403, detail="forbidden")
    docs = load_db()
    return {"count": len(docs)}


@app.get("/api/llm/status")
def api_llm_status(current_user=Depends(get_current_user)):
    if not getattr(current_user, "is_admin", False):
        raise HTTPException(status_code=403, detail="forbidden")
    return {
        "minimax": {
            "has_key": bool(os.getenv("MINIMAX_API_KEY")),
            "model": os.getenv("MINIMAX_MODEL"),
            "base_url": os.getenv("MINIMAX_BASE_URL"),
            "embed_model": os.getenv("MINIMAX_EMBED_MODEL")
        },
        "anthropic": {
            "has_key": bool(os.getenv("ANTHROPIC_API_KEY")),
            "model": os.getenv("ANTHROPIC_MODEL"),
            "base_url": os.getenv("ANTHROPIC_BASE_URL")
        }
    }


@app.get("/", response_class=HTMLResponse)
def frontend_index():
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    html_path = os.path.join(base_dir, "frontend", "demo.html")
    if not os.path.isfile(html_path):
        raise HTTPException(status_code=404, detail="frontend not found")
    with open(html_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


class ImportExcelPayload(BaseModel):
    path: Optional[str] = None
    move_to_knowledge: Optional[bool] = True


@app.post("/api/universities/import_excel")
def import_excel(payload: ImportExcelPayload, current_user=Depends(get_current_user)):
    if not getattr(current_user, "is_admin", False):
        raise HTTPException(status_code=403, detail="forbidden")
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    src_path = payload.path or os.path.join(base_dir, "Hackathon_Data sheet (1).xlsx")
    if not os.path.isfile(src_path):
        raise HTTPException(status_code=404, detail="excel not found")
    try:
        import openpyxl
    except Exception as e:
        raise HTTPException(status_code=500, detail="openpyxl not installed")
    try:
        wb = openpyxl.load_workbook(filename=src_path, data_only=True, read_only=True)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    rows = []
    for ws in wb.worksheets:
        header = None
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                header = [str(c).strip().lower() if c is not None else "" for c in row]
                continue
            rec = {header[j]: row[j] for j in range(len(header))}
            rows.append(rec)
    def g(v):
        return v if v is not None and str(v).strip() != "" else None
    def to_int(v):
        try:
            if v is None or str(v).strip() == "":
                return None
            return int(float(v))
        except Exception:
            return None
    def to_float(v):
        try:
            if v is None or str(v).strip() == "":
                return None
            return float(v)
        except Exception:
            return None
    def split_list(v):
        if v is None:
            return []
        if isinstance(v, list):
            return [str(x).strip() for x in v if x]
        s = str(v)
        if not s.strip():
            return []
        parts = [p.strip() for p in s.replace("；",";").replace(",",";").split(";")]
        return [p for p in parts if p]
    def slug(s):
        return "".join(ch.lower() if str(ch).isalnum() else "_" for ch in str(s)).strip("_")
    db = SessionLocal()
    imported = 0
    try:
        for r in rows:
            name = g(r.get("university") or r.get("name"))
            city = g(r.get("city"))
            if not name:
                continue
            qs_rank = to_int(r.get("qs_rank"))
            cn_rank = to_int(r.get("cn_rank"))
            fees = to_int(r.get("fees"))
            deadline = g(r.get("deadline"))
            fields = split_list(r.get("fields"))
            scholarships = split_list(r.get("scholarships"))
            official_link = g(r.get("official_link"))
            gpa_min = to_float(r.get("gpa_min"))
            toefl_min = to_int(r.get("toefl_min"))
            ielts_min = to_float(r.get("ielts_min"))
            hsk_min_level = to_int(r.get("hsk_min_level"))
            hsk_min_score = to_int(r.get("hsk_min_score"))
            special_requirements = g(r.get("special_requirements"))
            ext_id = f"{slug(name)}_{slug(city or 'cn')}"
            req = {
                "gpa_min": gpa_min,
                "toefl_min": toefl_min,
                "ielts_min": ielts_min,
                "hsk_required": bool(hsk_min_level or hsk_min_score),
                "hsk_min_level": hsk_min_level,
                "hsk_min_score": hsk_min_score,
                "special_requirements": special_requirements
            }
            u = db.query(University).filter(University.ext_id == ext_id).first()
            if not u:
                u = University(
                    ext_id=ext_id,
                    name=name,
                    city=city,
                    qs_rank=qs_rank,
                    cn_rank=cn_rank,
                    fees=fees,
                    deadline=deadline,
                    scholarships=json.dumps(scholarships, ensure_ascii=False),
                    official_link=official_link,
                    fields_offered=json.dumps(fields, ensure_ascii=False),
                    requirements=json.dumps(req, ensure_ascii=False)
                )
                db.add(u)
            else:
                u.name = name
                u.city = city
                u.qs_rank = qs_rank
                u.cn_rank = cn_rank
                u.fees = fees
                u.deadline = deadline
                u.scholarships = json.dumps(scholarships, ensure_ascii=False)
                u.official_link = official_link
                u.fields_offered = json.dumps(fields, ensure_ascii=False)
                u.requirements = json.dumps(req, ensure_ascii=False)
            imported += 1
        db.commit()
    finally:
        db.close()
    moved = None
    if payload.move_to_knowledge:
        try:
            knowledge_dir = os.path.join(base_dir, "knowledge")
            os.makedirs(knowledge_dir, exist_ok=True)
            target = os.path.join(knowledge_dir, os.path.basename(src_path))
            if os.path.abspath(src_path) != os.path.abspath(target):
                try:
                    os.replace(src_path, target)
                except Exception:
                    import shutil
                    shutil.copy2(src_path, target)
                    try:
                        os.remove(src_path)
                    except Exception:
                        pass
            moved = target
        except Exception:
            moved = None
    return {"imported": imported, "moved_to": moved}
